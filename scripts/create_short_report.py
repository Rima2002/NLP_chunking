from __future__ import annotations

from pathlib import Path
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")


PROJECT_ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = PROJECT_ROOT.parent
OUTPUT_PATH = REPORT_DIR / "NLP_rapor_kisa.docx"
GRAPH_PATH = PROJECT_ROOT / "outputs" / "all_confusion_matrices.png"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_text(cell, header, bold=True)
        set_cell_shading(cell, "D9EAF7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(3)


def add_body(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(4)


def add_compact_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.add_run(item)


def set_document_style(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.35)
    section.bottom_margin = Cm(1.35)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)

    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    for style_name in ["Heading 1", "Heading 2"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.bold = True
    styles["Heading 1"].font.size = Pt(13)
    styles["Heading 2"].font.size = Pt(11.5)


def add_cover(document: Document) -> None:
    for line in [
        "T.C.",
        "BURSA TEKNİK ÜNİVERSİTESİ",
        "BİLGİSAYAR MÜHENDİSLİĞİ BÖLÜMÜ",
    ]:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.bold = True
        r.font.size = Pt(13)

    document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("İsim ve Diğer Öbeklerin Saptanması (Chunking)")
    run.bold = True
    run.font.size = Pt(16)

    document.add_paragraph()
    info_rows = [
        ["Ders Adı", "Doğal Dil İşlemeye Giriş"],
        ["Programlama Dili", "Python"],
        ["Yöntem", "Conditional Random Fields (CRF)"],
        ["Github", "https://github.com/Rima2002/NLP_chunking"],
        ["Adı Soyadı", "Rima Farah Eleuch"],
        ["Öğrenci Numarası", "21360859216"],
    ]
    add_table(document, ["Bilgi", "Açıklama"], info_rows)

    document.add_paragraph()
    add_heading(document, "Kısa Özet", level=1)
    add_body(
        document,
        "Bu projede Türkçe cümlelerde isim öbeği, eylem öbeği, zarf öbeği ve "
        "cümlecik yapılarının otomatik olarak etiketlenmesi amaçlanmıştır. "
        "Nested chunking yaklaşımıyla her token için CHUNK-OUTER, CHUNK-INNER "
        "ve CLAUSE olmak üzere üç ayrı etiket sütunu tahmin edilmiştir.",
    )
    add_body(
        document,
        "Rapor kısa tutulmuş; sonuç bölümünde özellikle accuracy tablosu, "
        "support değerleri ve confusion matrix grafiklerinin yorumu öne çıkarılmıştır.",
    )


def add_dataset_and_method(document: Document) -> None:
    document.add_page_break()
    add_heading(document, "1. Veri Seti ve İşaretleme", level=1)
    add_body(
        document,
        "Veriler CoNLL formatında hazırlanmıştır. Her satır bir tokeni, boş satırlar "
        "ise cümle sınırını gösterir. Kullanılan temel biçim: ID FORM CHUNK-OUTER "
        "CHUNK-INNER CLAUSE.",
    )
    add_table(
        document,
        ["Veri Bölümü", "Cümle", "Token", "Amaç"],
        [
            ["Eğitim", "320", "2173", "CRF modellerini öğrenmek"],
            ["Test", "80", "543", "Model başarısını ölçmek"],
            ["Toplam", "400", "2716", "Proje veri seti"],
        ],
    )
    add_body(
        document,
        "CHUNK-OUTER sütunu dış öbekleri gösterir: NP isim öbeği, VP eylem öbeği, "
        "ADVP zarf öbeği ve O dışarıda kalan token anlamındadır. CHUNK-INNER "
        "sütunu iç içe geçen ilgi cümleciklerini RELCL etiketiyle gösterir. CLAUSE "
        "sütunu ise RELCL ve COMPCL gibi cümlecik yapılarını belirtir.",
    )

    add_heading(document, "2. Yöntem", level=1)
    add_body(
        document,
        "Modelleme için sıralı etiketleme problemlerinde kullanılan Conditional "
        "Random Fields (CRF) yöntemi seçilmiştir. Tek model yerine üç ayrı model "
        "eğitilmiştir: OUTER modeli dış öbekleri, INNER modeli iç yapıları, CLAUSE "
        "modeli ise cümlecikleri tahmin eder.",
    )
    add_compact_bullets(
        document,
        [
            "train.py: data/train.conll dosyasını kullanarak üç CRF modelini eğitir.",
            "evaluate.py: data/test.conll üzerinde precision, recall, f-measure, support ve accuracy hesaplar.",
            "predict.py: Kullanıcının girdiği yeni cümle için üç sütunlu nested chunking tahmini üretir.",
        ],
    )


def add_results(document: Document) -> None:
    document.add_page_break()
    add_heading(document, "3. Değerlendirme Sonuçları", level=1)
    add_body(
        document,
        "Değerlendirme eğitim verisi üzerinde değil, modelin eğitimde görmediği "
        "80 test cümlesi üzerinde yapılmıştır. Test bölümünde toplam 543 token vardır.",
    )
    add_table(
        document,
        ["Hedef Sütun", "Accuracy", "Kısa Yorum"],
        [
            ["CHUNK-OUTER", "0.9982", "Dış öbek etiketlerinde yalnızca çok az karışma vardır."],
            ["CHUNK-INNER", "1.0000", "RELCL ve boş iç yapı etiketleri hatasız tahmin edilmiştir."],
            ["CLAUSE", "1.0000", "RELCL, COMPCL ve O sınıfları testte doğru ayrılmıştır."],
        ],
    )

    add_heading(document, "Metriklerin Anlamı", level=2)
    add_compact_bullets(
        document,
        [
            "Precision: Modelin bir etiketi verdiğinde ne kadar doğru olduğunu gösterir.",
            "Recall: Gerçekte o etikete ait tokenlerin ne kadarının yakalandığını gösterir.",
            "F-measure: Precision ve recall değerlerinin dengeli ortalamasıdır.",
            "Support: Test verisinde ilgili sınıfa ait gerçek token sayısıdır.",
        ],
    )
    add_body(
        document,
        "Örneğin CLAUSE raporunda B-COMPCL için support 19’dur; bu, test verisinde "
        "gerçek etiketi B-COMPCL olan 19 token bulunduğu anlamına gelir. CLAUSE "
        "tablosundaki support değerleri 19 + 33 + 25 + 47 + 419 = 543 eder ve bu "
        "toplam test token sayısına eşittir.",
    )

    add_heading(document, "Sonuçların Yorumlanması", level=2)
    add_body(
        document,
        "CHUNK-INNER ve CLAUSE sonuçlarının 1.0000 çıkması, test verisindeki "
        "cümlecik etiketlerinin model tarafından hatasız tahmin edildiğini gösterir. "
        "Bu yüksek başarı yorumlanırken veri setinin proje kapsamında şablonlara "
        "dayalı olarak genişletildiği dikkate alınmalıdır. Gerçek ve daha çeşitli "
        "bir corpus üzerinde sonuçların daha düşük çıkması mümkündür.",
    )


def add_graph_page(document: Document) -> None:
    document.add_page_break()
    add_heading(document, "4. Grafiklerin Açıklanması", level=1)
    add_body(
        document,
        "Aşağıdaki birleşik confusion matrix görseli üç hedef sütun için gerçek "
        "etiketlerle model tahminlerini karşılaştırır. Satırlar gerçek sınıfları, "
        "sütunlar tahmin edilen sınıfları gösterir. Ana köşegen üzerindeki değerler "
        "doğru tahminleri, köşegen dışındaki değerler hataları ifade eder.",
    )
    if GRAPH_PATH.exists():
        picture = document.add_picture(str(GRAPH_PATH), width=Inches(6.95))
        picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption = document.add_paragraph("Şekil 1. OUTER, INNER ve CLAUSE modelleri için confusion matrix grafikleri")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].italic = True
    caption.runs[0].font.size = Pt(9)

    add_compact_bullets(
        document,
        [
            "OUTER grafiğinde değerlerin neredeyse tamamı köşegendedir; yalnızca B-ADVP sınıfında 1 token B-NP olarak karışmıştır.",
            "INNER grafiğinde B-RELCL, I-RELCL ve _ sınıflarının tamamı köşegendedir; bu yüzden accuracy 1.0000 olmuştur.",
            "CLAUSE grafiğinde B-COMPCL, B-RELCL, I-COMPCL, I-RELCL ve O sınıfları tamamen doğru ayrılmıştır.",
        ],
    )


def add_conclusion(document: Document) -> None:
    document.add_page_break()
    add_heading(document, "5. Tahmin ve Sonuç", level=1)
    add_body(
        document,
        "predict.py dosyası kullanıcıdan alınan yeni Türkçe cümleyi tokenlere ayırır "
        "ve her token için üç farklı çıktı üretir: CHUNK-OUTER, CHUNK-INNER ve CLAUSE. "
        "Bu yapı sayesinde yalnızca dış öbekler değil, iç içe geçen ilgi cümlecikleri "
        "ve tümleç cümlecikleri de ayrıca gösterilebilir.",
    )
    add_body(
        document,
        "Sonuç olarak proje, Türkçe nested chunking problemini CoNLL formatında "
        "hazırlanan veri üzerinde CRF yöntemiyle çözmüştür. 400 cümlelik veri seti "
        "320 eğitim ve 80 test cümlesi olarak ayrılmış, değerlendirme test verisi "
        "üzerinde yapılmıştır. Grafikler, modelin özellikle cümlecik etiketlerini "
        "hatasız ayırdığını göstermektedir.",
    )
    add_body(
        document,
        "Bununla birlikte başarı oranlarının çok yüksek çıkmasının nedeni, veri "
        "setinin proje kapsamında düzenli ve şablonlu örneklerle genişletilmiş "
        "olmasıdır. Daha gerçekçi bir değerlendirme için ileride farklı kaynaklardan "
        "toplanan daha çeşitli Türkçe cümleler elle etiketlenerek test edilebilir.",
    )
    add_heading(document, "Kullanılan Çıktı Dosyaları", level=2)
    add_compact_bullets(
        document,
        [
            "outputs/results_summary.txt: Genel metrik özeti.",
            "outputs/all_confusion_matrices.png: Birleşik confusion matrix grafiği.",
            "outputs/nested_predictions.conll: Gerçek ve tahmin etiketlerini içeren çıktı dosyası.",
        ],
    )


def main() -> None:
    document = Document()
    set_document_style(document)
    add_cover(document)
    add_dataset_and_method(document)
    add_results(document)
    add_graph_page(document)
    add_conclusion(document)
    document.save(OUTPUT_PATH)
    print(f"Kısa rapor oluşturuldu: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
