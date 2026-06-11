from __future__ import annotations

from pathlib import Path
import sys

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")


PROJECT_ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = PROJECT_ROOT.parent
SCREENSHOT_DIR = REPORT_DIR / "Screenshots"
OUTPUT_PATH = REPORT_DIR / "NLP_rapor_duzeltilmis.docx"

TRAIN_SCREENSHOT = SCREENSHOT_DIR / "1 train.png"
PREDICT_SCREENSHOT = SCREENSHOT_DIR / "3 predict.png"
CONFUSION_GRAPH = PROJECT_ROOT / "outputs" / "all_confusion_matrices.png"


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def write_cell(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(8.8)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, header in enumerate(headers):
        write_cell(table.rows[0].cells[index], header, bold=True)
        shade_cell(table.rows[0].cells[index], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            write_cell(cells[index], value)


def heading(document: Document, text: str, level: int = 1) -> None:
    p = document.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)


def body(document: Document, text: str) -> None:
    p = document.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)


def bullets(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        p.add_run(item)


def caption(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(8.8)


def add_picture_if_exists(document: Document, path: Path, width: float) -> None:
    if path.exists():
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Inches(width))


def set_style(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.45)
    section.right_margin = Cm(1.45)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.2)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    for style_name in ["Heading 1", "Heading 2"]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    document.styles["Heading 1"].font.size = Pt(13)
    document.styles["Heading 2"].font.size = Pt(11.3)


def cover_and_intro(document: Document) -> None:
    for line in ["T.C.", "BURSA TEKNİK ÜNİVERSİTESİ", "BİLGİSAYAR MÜHENDİSLİĞİ BÖLÜMÜ"]:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.bold = True
        r.font.size = Pt(12.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("İsim ve Diğer Öbeklerin Saptanması (Chunking)")
    r.bold = True
    r.font.size = Pt(15)

    add_table(
        document,
        ["Bilgi", "Açıklama"],
        [
            ["Ders", "Doğal Dil İşlemeye Giriş"],
            ["Programlama Dili", "Python"],
            ["Yöntem", "Conditional Random Fields (CRF)"],
            ["Github", "https://github.com/Rima2002/NLP_chunking"],
            ["Öğrenci", "Rima Farah Eleuch - 21360859216"],
        ],
    )

    heading(document, "1. Giriş")
    body(
        document,
        "Bu projede Türkçe cümlelerdeki isim öbeği, eylem öbeği, zarf öbeği ve "
        "cümlecik yapıları otomatik olarak etiketlenmiştir. Çalışmada nested "
        "chunking yapısı kullanılmıştır; böylece her token için CHUNK-OUTER, "
        "CHUNK-INNER ve CLAUSE olmak üzere üç ayrı etiket üretilmiştir.",
    )
    body(
        document,
        "Rapor, son kod çıktılarıyla uyumlu olacak şekilde kısa tutulmuştur. "
        "Özellikle eğitim/test veri sayısı, güncel başarı değerleri ve confusion "
        "matrix grafiklerinin yorumu verilmiştir.",
    )


def data_and_method(document: Document) -> None:
    heading(document, "2. Veri Seti ve Yöntem")
    body(
        document,
        "Veri seti CoNLL formatındadır. Her satır bir tokeni gösterir; boş satırlar "
        "cümle sınırıdır. Kullanılan sütunlar: ID, FORM, CHUNK-OUTER, CHUNK-INNER "
        "ve CLAUSE.",
    )
    add_table(
        document,
        ["Bölüm", "Cümle Örneği", "Token", "Kullanım"],
        [
            ["Eğitim", "320", "2173", "train.py ile model eğitimi"],
            ["Test", "80", "543", "evaluate.py ile değerlendirme"],
            ["Toplam", "400", "2716", "Tüm proje verisi"],
        ],
    )
    body(
        document,
        "Eğitim işlemi data/train.conll dosyasıyla, değerlendirme ise data/test.conll "
        "dosyasıyla yapılmıştır. Bu nedenle evaluate.py çıktıları eğitim verisini "
        "değil, modelin eğitimde görmediği 80 test cümlesini göstermektedir.",
    )
    bullets(
        document,
        [
            "CHUNK-OUTER: NP, VP, ADVP ve O gibi dış öbek etiketlerini gösterir.",
            "CHUNK-INNER: İç içe geçen ilgi cümlecikleri için RELCL etiketini kullanır.",
            "CLAUSE: RELCL ve COMPCL gibi cümlecik yapılarını gösterir.",
        ],
    )
    body(
        document,
        "Model olarak CRF seçilmiştir. OUTER, INNER ve CLAUSE için üç ayrı CRF modeli "
        "eğitilmiş ve çıktılar outputs klasörüne kaydedilmiştir.",
    )


def training_and_results(document: Document) -> None:
    document.add_page_break()
    heading(document, "3. Eğitim ve Test Sonuçları")
    add_picture_if_exists(document, TRAIN_SCREENSHOT, 5.9)
    caption(document, "Şekil 1. train.py çıktısı: OUTER, INNER ve CLAUSE modellerinin eğitilmesi")

    add_table(
        document,
        ["Model", "Accuracy", "Sonuç"],
        [
            ["CHUNK-OUTER", "0.9982", "Dış öbeklerde yalnızca çok küçük bir karışma vardır."],
            ["CHUNK-INNER", "1.0000", "İç yapı etiketleri test verisinde hatasızdır."],
            ["CLAUSE", "1.0000", "Cümlecik etiketleri test verisinde hatasızdır."],
        ],
    )
    body(
        document,
        "Support değeri, test verisinde ilgili sınıfa ait gerçek token sayısıdır. "
        "Örneğin CLAUSE çıktısında B-COMPCL support değeri 19, B-RELCL support değeri "
        "33, I-COMPCL support değeri 25, I-RELCL support değeri 47 ve O support değeri "
        "419’dur. Bu değerlerin toplamı 543’tür ve test token sayısına eşittir.",
    )
    body(
        document,
        "CHUNK-INNER ve CLAUSE için tüm precision, recall ve f-measure değerlerinin "
        "1.0000 çıkması, test verisindeki ilgili etiketlerin tamamının doğru tahmin "
        "edildiğini gösterir. Bu sonuç yorumlanırken veri setinin proje kapsamında "
        "düzenli örneklerle genişletildiği dikkate alınmalıdır.",
    )


def graph_explanation(document: Document) -> None:
    document.add_page_break()
    heading(document, "4. Grafiklerin Açıklanması")
    body(
        document,
        "Confusion matrix grafiklerinde satırlar gerçek etiketleri, sütunlar model "
        "tahminlerini gösterir. Ana köşegen üzerindeki değerler doğru tahminlerdir. "
        "Köşegen dışındaki değerler ise modelin hangi sınıfları karıştırdığını gösterir.",
    )
    add_picture_if_exists(document, CONFUSION_GRAPH, 6.95)
    caption(document, "Şekil 2. Güncel OUTER, INNER ve CLAUSE confusion matrix grafikleri")
    bullets(
        document,
        [
            "OUTER grafiğinde değerlerin neredeyse tamamı köşegendedir; yalnızca B-ADVP sınıfından 1 token B-NP olarak tahmin edilmiştir.",
            "INNER grafiğinde B-RELCL, I-RELCL ve _ sınıflarının tamamı doğru tahmin edilmiştir; bu yüzden accuracy 1.0000’dır.",
            "CLAUSE grafiğinde B-COMPCL, B-RELCL, I-COMPCL, I-RELCL ve O sınıflarının tamamı köşegendedir; model test verisinde cümlecikleri hatasız ayırmıştır.",
        ],
    )


def prediction_and_conclusion(document: Document) -> None:
    document.add_page_break()
    heading(document, "5. Tahmin ve Sonuç")
    add_picture_if_exists(document, PREDICT_SCREENSHOT, 5.8)
    caption(document, "Şekil 3. predict.py ile örnek cümle üzerinde nested chunking tahmini")
    body(
        document,
        "predict.py dosyası kullanıcıdan alınan Türkçe cümleyi tokenlere ayırır ve "
        "her token için üç ayrı etiket üretir. Örnek çıktıda model dış öbekleri "
        "CHUNK-OUTER sütununda, iç yapıları CHUNK-INNER sütununda, cümlecik bilgisini "
        "ise CLAUSE sütununda göstermektedir.",
    )
    body(
        document,
        "Sonuç olarak proje, 400 cümlelik veri seti üzerinde CRF tabanlı nested "
        "chunking uygulamasını gerçekleştirmiştir. Eğitim 320 cümleyle, test ise "
        "80 cümleyle yapılmıştır. Güncel grafikler ve metrikler modelin özellikle "
        "INNER ve CLAUSE etiketlerinde test verisini hatasız tahmin ettiğini "
        "göstermektedir.",
    )
    body(
        document,
        "Başarı oranlarının çok yüksek çıkmasının nedeni, veri setinin düzenli ve "
        "şablonlu örneklerle genişletilmiş olmasıdır. Daha genel bir değerlendirme "
        "için farklı kaynaklardan toplanmış daha çeşitli Türkçe cümleler kullanılabilir.",
    )


def main() -> None:
    document = Document()
    set_style(document)
    cover_and_intro(document)
    data_and_method(document)
    training_and_results(document)
    graph_explanation(document)
    prediction_and_conclusion(document)
    document.save(OUTPUT_PATH)
    print(f"Düzeltilmiş rapor oluşturuldu: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
